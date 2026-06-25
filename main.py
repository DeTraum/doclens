"""
DocLens MVP v0.2.0 - 双模式：本地上传 + 飞书链接
启动: python main.py
访问: http://localhost:8910
"""

import os
import sys
import uuid
import time
import json
import shutil
import socket
import asyncio
import logging
import subprocess
import re
import httpx
from pathlib import Path
from datetime import datetime
from contextlib import asynccontextmanager

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Body
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse, RedirectResponse
from fastapi.middleware.cors import CORSMiddleware
import uvicorn

# ============================================================
# 配置
# ============================================================

BASE_DIR = Path(__file__).parent
UPLOAD_DIR = BASE_DIR / "data" / "uploads"
OUTPUT_DIR = BASE_DIR / "data" / "outputs"
STATIC_DIR = BASE_DIR / "static"
LOG_DIR = BASE_DIR / "data" / "logs"
CONFIG_FILE = BASE_DIR / "config.json"

for d in [UPLOAD_DIR, OUTPUT_DIR, LOG_DIR]:
    d.mkdir(parents=True, exist_ok=True)

HOST = "127.0.0.1"
PORT = 8910

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler(LOG_DIR / "doclens.log", encoding="utf-8"),
    ],
)
logger = logging.getLogger("doclens")


def load_config() -> dict:
    """加载飞书配置"""
    if CONFIG_FILE.exists():
        with open(CONFIG_FILE, "r") as f:
            return json.load(f)
    return {"feishu_app_id": "", "feishu_app_secret": ""}


def save_config(cfg: dict):
    with open(CONFIG_FILE, "w") as f:
        json.dump(cfg, f, indent=2, ensure_ascii=False)


# ============================================================
# 网络控制
# ============================================================

class NetworkGuard:
    """网络状态检测与管理"""

    # 多个探测目标(国内可靠的 DNS + 飞书)，任意一个通即视为在线
    PROBES = [("223.5.5.5", 53), ("119.29.29.29", 53), ("open.feishu.cn", 443), ("8.8.8.8", 53)]

    @staticmethod
    def is_online() -> bool:
        for host, port in NetworkGuard.PROBES:
            try:
                s = socket.create_connection((host, port), timeout=2)
                s.close()
                return True
            except (socket.timeout, OSError):
                continue
        return False

    @staticmethod
    def check_feishu_reachable() -> bool:
        try:
            s = socket.create_connection(("open.feishu.cn", 443), timeout=3)
            s.close()
            return True
        except (socket.timeout, OSError):
            return False

    @staticmethod
    def get_status() -> dict:
        feishu = NetworkGuard.check_feishu_reachable()
        online = feishu or NetworkGuard.is_online()
        return {
            "online": online,
            "feishu_reachable": feishu,
            "recommendation": "safe" if not online else "warn",
        }


network_guard = NetworkGuard()

# ============================================================
# 飞书 API 客户端
# ============================================================

class FeishuClient:
    """飞书开放平台 API 封装"""

    BASE_URL = "https://open.feishu.cn/open-apis"

    # 用户登录(OAuth)需要的权限范围
    USER_SCOPES = "docx:document drive:drive bitable:app offline_access"

    def __init__(self):
        self.app_id = ""
        self.app_secret = ""
        self.tenant_token = ""
        self.token_expires = 0
        # 用户身份令牌(登录后才有)
        self.user_token = ""
        self.user_refresh_token = ""
        self.user_token_expires = 0
        self.user_name = ""
        self.active_token = ""  # _headers() 实际用的令牌

    def configure(self, app_id: str, app_secret: str):
        self.app_id = app_id
        self.app_secret = app_secret
        self.tenant_token = ""
        self.token_expires = 0
        # 从 config 恢复用户登录态
        cfg = load_config()
        self.user_refresh_token = cfg.get("user_refresh_token", "")
        self.user_name = cfg.get("user_name", "")
        self.user_token = ""
        self.user_token_expires = 0

    @property
    def is_configured(self) -> bool:
        return bool(self.app_id and self.app_secret)

    @property
    def has_user_auth(self) -> bool:
        return bool(self.user_refresh_token)

    def authorize_url(self, redirect_uri: str, state: str = "doclens") -> str:
        from urllib.parse import urlencode
        q = urlencode({
            "client_id": self.app_id,
            "redirect_uri": redirect_uri,
            "scope": self.USER_SCOPES,
            "state": state,
            "response_type": "code",
        })
        return f"https://accounts.feishu.cn/open-apis/authen/v1/authorize?{q}"

    async def exchange_code(self, code: str, redirect_uri: str):
        """用授权码换取用户令牌，保存到 config"""
        async with httpx.AsyncClient() as client:
            r = await client.post(
                f"{self.BASE_URL}/authen/v2/oauth/token",
                json={"grant_type": "authorization_code", "client_id": self.app_id,
                      "client_secret": self.app_secret, "code": code,
                      "redirect_uri": redirect_uri},
            )
            d = r.json()
        if "access_token" not in d:
            raise RuntimeError(f"换取令牌失败: {d.get('error_description') or d.get('msg') or d}")
        self.user_token = d["access_token"]
        self.user_token_expires = time.time() + d.get("expires_in", 7200) - 60
        self.user_refresh_token = d.get("refresh_token", self.user_refresh_token)
        # 取用户名
        try:
            async with httpx.AsyncClient() as client:
                ui = await client.get(f"{self.BASE_URL}/authen/v1/user_info",
                                      headers={"Authorization": f"Bearer {self.user_token}"})
                self.user_name = ui.json().get("data", {}).get("name", "")
        except Exception:
            pass
        cfg = load_config()
        cfg["user_refresh_token"] = self.user_refresh_token
        cfg["user_name"] = self.user_name
        save_config(cfg)

    async def _ensure_user_token(self):
        """刷新用户 access_token"""
        if self.user_token and time.time() < self.user_token_expires:
            return True
        if not self.user_refresh_token:
            return False
        async with httpx.AsyncClient() as client:
            r = await client.post(
                f"{self.BASE_URL}/authen/v2/oauth/token",
                json={"grant_type": "refresh_token", "client_id": self.app_id,
                      "client_secret": self.app_secret,
                      "refresh_token": self.user_refresh_token},
            )
            d = r.json()
        if "access_token" not in d:
            logger.warning(f"用户令牌刷新失败，需重新登录: {d.get('error_description') or d.get('msg')}")
            self.user_refresh_token = ""
            cfg = load_config(); cfg["user_refresh_token"] = ""; save_config(cfg)
            return False
        self.user_token = d["access_token"]
        self.user_token_expires = time.time() + d.get("expires_in", 7200) - 60
        if d.get("refresh_token"):
            self.user_refresh_token = d["refresh_token"]
            cfg = load_config(); cfg["user_refresh_token"] = self.user_refresh_token; save_config(cfg)
        return True

    def logout(self):
        self.user_token = self.user_refresh_token = self.user_name = ""
        self.user_token_expires = 0
        cfg = load_config()
        cfg["user_refresh_token"] = ""; cfg["user_name"] = ""
        save_config(cfg)

    async def _ensure_token(self):
        """确保有可用令牌。优先用户身份(产出归你所有)，否则回退应用身份。"""
        if self.user_refresh_token and await self._ensure_user_token():
            self.active_token = self.user_token
            return
        self.active_token = await self._ensure_tenant_token()

    async def _ensure_tenant_token(self) -> str:
        """获取/刷新应用身份令牌(始终可用，与用户登录无关)"""
        if not (self.tenant_token and time.time() < self.token_expires):
            async with httpx.AsyncClient() as client:
                resp = await client.post(
                    f"{self.BASE_URL}/auth/v3/tenant_access_token/internal",
                    json={"app_id": self.app_id, "app_secret": self.app_secret},
                )
                data = resp.json()
                if data.get("code") != 0:
                    raise RuntimeError(f"飞书鉴权失败: {data.get('msg', '未知错误')}")
                self.tenant_token = data["tenant_access_token"]
                self.token_expires = time.time() + data.get("expire", 7200) - 60
        return self.tenant_token

    @property
    def using_user_identity(self) -> bool:
        return bool(self.active_token and self.active_token == self.user_token and self.user_token)

    def _headers(self) -> dict:
        return {"Authorization": f"Bearer {self.active_token}"}

    def _tenant_headers(self) -> dict:
        """管理表(多维表格)固定用应用身份读写，避免随登录状态切换导致记录分散到多张表"""
        return {"Authorization": f"Bearer {self.tenant_token}"}

    def _user_headers(self) -> dict:
        return {"Authorization": f"Bearer {self.user_token}"}

    INDEX_NAME = "DocLens 转换记录"

    async def _search_user_bitable(self, name: str) -> str:
        """用用户身份在其云盘搜索同名多维表格，返回 app_token(找不到返回'')"""
        try:
            async with httpx.AsyncClient(timeout=30) as c:
                r = await c.post(f"{self.BASE_URL}/drive/v1/files/search",
                                 headers=self._user_headers(),
                                 json={"search_key": name, "count": 50, "docs_types": ["bitable"]})
                j = r.json()
                if j.get("code") != 0:
                    logger.warning(f"搜索管理表失败: {j.get('msg')}")
                    return ""
                data = j.get("data", {})
                # 兼容两种返回结构
                for f in data.get("files", []):
                    if f.get("name", "").strip() == name and f.get("type") == "bitable":
                        return f.get("token", "")
                for f in data.get("docs_entities", []):
                    if f.get("title", "").strip() == name and "bitable" in str(f.get("docs_type", "")):
                        return f.get("docs_token", "")
        except Exception as e:
            logger.warning(f"搜索管理表异常: {e}")
        return ""

    async def _first_table_id(self, app_token: str, user: bool) -> str:
        async with httpx.AsyncClient(timeout=30) as c:
            r = await c.get(f"{self.BASE_URL}/bitable/v1/apps/{app_token}/tables",
                            headers=self._user_headers() if user else self._tenant_headers())
            items = r.json().get("data", {}).get("items", [])
            return items[0]["table_id"] if items else ""

    def parse_doc_url(self, url: str) -> tuple[str, str]:
        """
        从飞书文档链接解析 doc_token 和文档类型
        支持格式:
          https://xxx.feishu.cn/docx/xxxToken
          https://xxx.feishu.cn/wiki/xxxToken
          https://xxx.feishu.cn/docs/xxxToken
        """
        patterns = [
            r"feishu\.cn/docx/([A-Za-z0-9]+)",
            r"feishu\.cn/wiki/([A-Za-z0-9]+)",
            r"feishu\.cn/docs/([A-Za-z0-9]+)",
            r"feishu\.cn/sheets/([A-Za-z0-9]+)",
            r"larksuite\.com/docx/([A-Za-z0-9]+)",
            r"larksuite\.com/wiki/([A-Za-z0-9]+)",
        ]
        for pat in patterns:
            m = re.search(pat, url)
            if m:
                doc_type = "docx" if "docx" in pat or "wiki" in pat else "doc"
                return m.group(1), doc_type

        raise ValueError("无法解析飞书文档链接，请检查 URL 格式")

    async def get_doc_title(self, doc_token: str) -> str:
        """获取源飞书文档的标题(取不到返回空)"""
        try:
            await self._ensure_token()
            async with httpx.AsyncClient(timeout=20) as client:
                r = await client.get(f"{self.BASE_URL}/docx/v1/documents/{doc_token}",
                                     headers=self._headers())
                return r.json().get("data", {}).get("document", {}).get("title", "") or ""
        except Exception:
            return ""

    async def get_doc_blocks(self, doc_token: str) -> list:
        """获取文档的所有 block (含图片 token)"""
        await self._ensure_token()
        blocks = []
        page_token = ""

        async with httpx.AsyncClient() as client:
            while True:
                params = {"document_id": doc_token, "page_size": 500}
                if page_token:
                    params["page_token"] = page_token

                resp = await client.get(
                    f"{self.BASE_URL}/docx/v1/documents/{doc_token}/blocks",
                    headers=self._headers(),
                    params=params,
                )
                data = resp.json()
                if data.get("code") != 0:
                    raise RuntimeError(f"获取文档内容失败: {data.get('msg')}")

                items = data.get("data", {}).get("items", [])
                blocks.extend(items)

                if not data.get("data", {}).get("has_more"):
                    break
                page_token = data["data"].get("page_token", "")

        return blocks

    async def extract_images_from_doc(self, doc_token: str) -> list[str]:
        """从文档中提取所有图片，下载到本地，返回本地路径列表"""
        await self._ensure_token()
        blocks = await self.get_doc_blocks(doc_token)

        image_tokens = []
        for block in blocks:
            if block.get("block_type") == 27:  # image block
                img = block.get("image", {})
                token = img.get("token")
                if token:
                    image_tokens.append(token)

        if not image_tokens:
            raise ValueError("文档中未找到图片")

        # 下载图片
        local_paths = []
        async with httpx.AsyncClient() as client:
            for idx, img_token in enumerate(image_tokens):
                resp = await client.get(
                    f"{self.BASE_URL}/drive/v1/medias/{img_token}/download",
                    headers=self._headers(),
                    follow_redirects=True,
                )
                if resp.status_code == 200:
                    ext = ".png"
                    ct = resp.headers.get("content-type", "")
                    if "jpeg" in ct or "jpg" in ct:
                        ext = ".jpg"
                    path = UPLOAD_DIR / f"feishu_{doc_token}_{idx}{ext}"
                    with open(path, "wb") as f:
                        f.write(resp.content)
                    local_paths.append(str(path))
                    logger.info(f"下载飞书图片 {idx+1}/{len(image_tokens)}")

        return local_paths

    # 飞书单次创建子块上限为 50，超过会报 field validation failed
    MAX_CHILDREN_PER_REQUEST = 50

    async def create_doc_with_content(self, title: str, markdown: str) -> str:
        """创建飞书文档并写入 Markdown 内容(含真实表格)，返回文档链接"""
        return await self.create_doc_from_items(title, self._parse_markdown(markdown))

    async def create_doc_from_items(self, title: str, items: list[tuple]) -> str:
        """创建飞书文档并写入有序条目(('block',dict)/('table',rows)/('image',path))"""
        await self._ensure_token()

        async with httpx.AsyncClient(timeout=60) as client:
            # 创建文档
            resp = await client.post(
                f"{self.BASE_URL}/docx/v1/documents",
                headers=self._headers(),
                json={"title": title},
            )
            data = resp.json()
            if data.get("code") != 0:
                raise RuntimeError(f"创建文档失败: {data.get('msg')}")

            doc_token = data["data"]["document"]["document_id"]

            # 获取文档第一个 block (page block) 的 ID
            resp2 = await client.get(
                f"{self.BASE_URL}/docx/v1/documents/{doc_token}/blocks",
                headers=self._headers(),
            )
            page_block_id = resp2.json()["data"]["items"][0]["block_id"]

            # 按顺序写入：普通块累积后分批(≤50)插入，表格用嵌套块 API 单独插入
            index = 0          # 当前已插入的顶层子块数(用于保持顺序)
            buffer: list[dict] = []

            async def flush_buffer():
                nonlocal index, buffer
                if not buffer:
                    return
                for i in range(0, len(buffer), self.MAX_CHILDREN_PER_REQUEST):
                    chunk = buffer[i:i + self.MAX_CHILDREN_PER_REQUEST]
                    r = await client.post(
                        f"{self.BASE_URL}/docx/v1/documents/{doc_token}/blocks/{page_block_id}/children",
                        headers=self._headers(),
                        json={"children": chunk, "index": index},
                    )
                    rj = r.json()
                    if rj.get("code") != 0:
                        logger.warning(f"写入文本块失败: {rj.get('msg')}")
                    else:
                        index += len(chunk)
                buffer = []

            for kind, payload in items:
                if kind == "table":
                    await flush_buffer()
                    ok = await self._insert_table(client, doc_token, page_block_id, payload, index)
                    if ok:
                        index += 1
                    else:
                        # 表格写入失败则降级为逐行文本，保证内容不丢
                        rows = {}
                        if isinstance(payload, dict):
                            for cell in payload.get("cells", []):
                                rows.setdefault(cell["row"], []).append(cell["text"])
                            lines = [" | ".join(rows[k]) for k in sorted(rows)]
                        else:
                            lines = [" | ".join(r) for r in payload]
                        for ln in lines:
                            buffer.append(self._text_block(ln))
                        await flush_buffer()
                elif kind == "image":
                    await flush_buffer()
                    ok = await self._insert_image(client, doc_token, page_block_id, payload, index)
                    if ok:
                        index += 1
                    else:
                        buffer.append(self._text_block("[图片插入失败]"))
                        await flush_buffer()
                else:
                    buffer.append(payload)
            await flush_buffer()

            # 用户身份创建的文档已归你所有(在你云盘、可编辑)，无需公开分享；
            # 仅当回退到机器人身份时，才开放编辑权限以便你能访问。
            if not self.using_user_identity:
                try:
                    await client.patch(
                        f"{self.BASE_URL}/drive/v1/permissions/{doc_token}/public?type=docx",
                        headers=self._headers(),
                        json={"link_share_entity": "anyone_editable"},
                    )
                except Exception as e:
                    logger.warning(f"开放编辑权限失败(可手动设置): {e}")

            return f"https://feishu.cn/docx/{doc_token}"

    # ---- 转换记录管理中心(多维表格) ----

    async def ensure_index_table(self):
        """返回 (app_token, table_id, use_user)。
        - 已登录：表绑定到你的飞书账号(在你云盘搜/建)，换设备登录同账号即同步同一张表
        - 未登录：用应用身份的本机表
        """
        cfg = load_config()
        if self.has_user_auth and await self._ensure_user_token():
            return await self._ensure_user_index(cfg)
        return await self._ensure_tenant_index(cfg)

    async def _ensure_user_index(self, cfg):
        """账号绑定的管理表：本机缓存 -> 云盘搜索 -> 创建。"""
        c = cfg.get("user_index", {})
        if c.get("app_token") and c.get("table_id"):
            return c["app_token"], c["table_id"], True
        # 换设备/首次：在用户云盘按名搜索已有的表
        app_token = await self._search_user_bitable(self.INDEX_NAME)
        table_id = await self._first_table_id(app_token, user=True) if app_token else ""
        index_url = ""
        if not (app_token and table_id):
            # 没找到 -> 在用户云盘新建(归你所有)
            async with httpx.AsyncClient(timeout=30) as client:
                r = await client.post(f"{self.BASE_URL}/bitable/v1/apps",
                                      headers=self._user_headers(), json={"name": self.INDEX_NAME})
                j = r.json()
                if j.get("code") != 0:
                    raise RuntimeError(f"创建管理表失败: {j.get('msg')}")
                app_token = j["data"]["app"]["app_token"]
                index_url = j["data"]["app"].get("url", "")
                body = {"table": {"name": "转换记录", "fields": [
                    {"field_name": "标题", "type": 1}, {"field_name": "链接", "type": 15},
                    {"field_name": "来源", "type": 1}, {"field_name": "时间", "type": 1}]}}
                r2 = await client.post(f"{self.BASE_URL}/bitable/v1/apps/{app_token}/tables",
                                       headers=self._user_headers(), json=body)
                table_id = r2.json()["data"]["table_id"]
        cfg["user_index"] = {"app_token": app_token, "table_id": table_id,
                             "url": index_url or f"https://feishu.cn/base/{app_token}"}
        cfg["index_url"] = cfg["user_index"]["url"]
        save_config(cfg)
        logger.info("账号管理表就绪(跨设备同步)")
        return app_token, table_id, True

    async def _ensure_tenant_index(self, cfg):
        """未登录时的本机应用表。"""
        if cfg.get("index_app_token") and cfg.get("index_table_id"):
            return cfg["index_app_token"], cfg["index_table_id"], False
        await self._ensure_tenant_token()
        async with httpx.AsyncClient(timeout=30) as client:
            r = await client.post(f"{self.BASE_URL}/bitable/v1/apps",
                                  headers=self._tenant_headers(), json={"name": self.INDEX_NAME})
            j = r.json()
            if j.get("code") != 0:
                raise RuntimeError(f"创建管理表失败: {j.get('msg')}")
            app_token = j["data"]["app"]["app_token"]
            index_url = j["data"]["app"].get("url", "")
            body = {"table": {"name": "转换记录", "fields": [
                {"field_name": "标题", "type": 1}, {"field_name": "链接", "type": 15},
                {"field_name": "来源", "type": 1}, {"field_name": "时间", "type": 1}]}}
            r2 = await client.post(f"{self.BASE_URL}/bitable/v1/apps/{app_token}/tables",
                                   headers=self._tenant_headers(), json=body)
            table_id = r2.json()["data"]["table_id"]
            try:
                await client.patch(f"{self.BASE_URL}/drive/v1/permissions/{app_token}/public?type=bitable",
                                   headers=self._tenant_headers(), json={"link_share_entity": "anyone_editable"})
            except Exception:
                pass
        cfg["index_app_token"] = app_token
        cfg["index_table_id"] = table_id
        cfg["index_url"] = index_url
        save_config(cfg)
        logger.info(f"管理表格已创建: {index_url}")
        return app_token, table_id, False

    async def append_index_record(self, title: str, doc_url: str, source: str):
        """把一次转换结果追加到管理表格(失败不影响主流程)。"""
        try:
            app_token, table_id, use_user = await self.ensure_index_table()
            headers = self._user_headers() if use_user else self._tenant_headers()
            rec = {"fields": {
                "标题": title,
                "链接": {"text": "打开文档", "link": doc_url},
                "来源": source,
                "时间": datetime.now().strftime("%Y-%m-%d %H:%M"),
            }}
            async with httpx.AsyncClient(timeout=30) as client:
                r = await client.post(
                    f"{self.BASE_URL}/bitable/v1/apps/{app_token}/tables/{table_id}/records",
                    headers=headers, json=rec)
                if r.json().get("code") != 0:
                    logger.warning(f"写入管理表格失败: {r.json().get('msg')}")
        except Exception as e:
            logger.warning(f"写入管理表格异常(不影响转换): {e}")

    async def _insert_table(self, client, doc_token, page_block_id, grid, index) -> bool:
        """用嵌套块 API 创建真实飞书表格，并还原合并单元格(rowspan/colspan)。
        兼容旧格式(grid 为二维数组)。"""
        # 兼容：若传入的是二维数组，转成 grid
        if isinstance(grid, list):
            n_rows = len(grid); n_cols = max((len(r) for r in grid), default=0)
            cells = [{"row": i, "col": j, "rs": 1, "cs": 1, "text": (grid[i][j] if j < len(grid[i]) else "")}
                     for i in range(n_rows) for j in range(n_cols)]
            grid = {"n_rows": n_rows, "n_cols": n_cols, "cells": cells}
        n_rows, n_cols = grid["n_rows"], grid["n_cols"]
        if n_rows <= 0 or n_cols <= 0:
            return True

        # 内容放到每个(合并)格的左上角；记录需要合并的区域
        content, merges = {}, []
        for cell in grid["cells"]:
            content[(cell["row"], cell["col"])] = cell["text"]
            if cell["rs"] > 1 or cell["cs"] > 1:
                merges.append(cell)

        descendants, cell_ids = [], []
        for i in range(n_rows):
            for j in range(n_cols):
                cid, tid = f"c_{i}_{j}", f"tx_{i}_{j}"
                cell_ids.append(cid)
                descendants.append({"block_id": cid, "block_type": 32,
                                    "table_cell": {}, "children": [tid]})
                descendants.append({"block_id": tid, "block_type": 2,
                                    "text": {"elements": [{"text_run": {"content": content.get((i, j)) or " "}}]}})
        table_block = {"block_id": "tbl", "block_type": 31,
                       "table": {"property": {"row_size": n_rows, "column_size": n_cols,
                                              "header_row": True}},
                       "children": cell_ids}
        descendants.insert(0, table_block)
        try:
            r = await client.post(
                f"{self.BASE_URL}/docx/v1/documents/{doc_token}/blocks/{page_block_id}/descendant",
                headers=self._headers(),
                json={"index": index, "children_id": ["tbl"], "descendants": descendants},
            )
            rj = r.json()
            if rj.get("code") != 0:
                logger.warning(f"创建表格失败，降级为文本: {rj.get('msg')}")
                return False
            # 取新建表格的真实 block_id，逐个合并单元格
            if merges:
                tbl_id = ""
                for b in rj.get("data", {}).get("children", []):
                    if b.get("block_type") == 31:
                        tbl_id = b.get("block_id", ""); break
                for m in merges:
                    if not tbl_id:
                        break
                    try:
                        await client.patch(
                            f"{self.BASE_URL}/docx/v1/documents/{doc_token}/blocks/{tbl_id}",
                            headers=self._headers(),
                            json={"merge_table_cells": {
                                "row_start_index": m["row"], "row_end_index": m["row"] + m["rs"],
                                "column_start_index": m["col"], "column_end_index": m["col"] + m["cs"]}},
                        )
                    except Exception as e:
                        logger.warning(f"合并单元格失败(不影响表格): {e}")
            return True
        except Exception as e:
            logger.warning(f"创建表格异常，降级为文本: {e}")
            return False

    async def _insert_image(self, client, doc_token, page_block_id, img_path, index) -> bool:
        """3 步插入图片：建空图片块 -> 上传媒体 -> 绑定 token"""
        try:
            # 1. 创建空图片块
            r = await client.post(
                f"{self.BASE_URL}/docx/v1/documents/{doc_token}/blocks/{page_block_id}/children",
                headers=self._headers(),
                json={"children": [{"block_type": 27, "image": {}}], "index": index},
            )
            rj = r.json()
            if rj.get("code") != 0:
                logger.warning(f"创建图片块失败: {rj.get('msg')}")
                return False
            img_block = rj["data"]["children"][0]["block_id"]

            # 2. 上传媒体(挂到图片块下)
            size = os.path.getsize(img_path)
            fname = os.path.basename(img_path)
            with open(img_path, "rb") as f:
                files = {"file": (fname, f, "application/octet-stream")}
                data = {"file_name": fname, "parent_type": "docx_image",
                        "parent_node": img_block, "size": str(size)}
                r2 = await client.post(
                    f"{self.BASE_URL}/drive/v1/medias/upload_all",
                    headers={"Authorization": self._headers()["Authorization"]},
                    data=data, files=files,
                )
            j2 = r2.json()
            if j2.get("code") != 0:
                logger.warning(f"上传图片失败: {j2.get('msg')}")
                return False
            token = j2["data"]["file_token"]

            # 3. 把 token 绑定到图片块
            r3 = await client.patch(
                f"{self.BASE_URL}/docx/v1/documents/{doc_token}/blocks/{img_block}",
                headers=self._headers(),
                json={"replace_image": {"token": token}},
            )
            if r3.json().get("code") != 0:
                logger.warning(f"绑定图片失败: {r3.json().get('msg')}")
                return False
            return True
        except Exception as e:
            logger.warning(f"插入图片异常: {e}")
            return False

    def _parse_markdown(self, markdown: str) -> list[tuple]:
        """Markdown -> 有序条目列表。元素为 ('block', dict) 或 ('table', rows)。"""
        items: list[tuple] = []
        # 先按 <table>...</table> 切分(MinerU 表格为单行 HTML)
        parts = re.split(r"(<table>.*?</table>)", markdown, flags=re.DOTALL)
        for part in parts:
            if not part.strip():
                continue
            if part.strip().startswith("<table>"):
                grid = self._parse_html_table_grid(part)
                if grid and grid["cells"]:
                    items.append(("table", grid))
                continue
            for line in part.split("\n"):
                line = line.strip()
                if not line:
                    continue
                # 过滤 HTML 注释(如 <!-- 第 1 页 -->)；--- 交给下面转成分割线
                if re.fullmatch(r"<!--.*-->", line):
                    continue
                # 图片行 ![](path)：解析出本地文件，作为独立图片条目
                img_match = re.search(r"!\[[^\]]*\]\(([^)]+)\)", line)
                if img_match:
                    local = self._resolve_image_path(img_match.group(1))
                    if local:
                        items.append(("image", local))
                        continue
                    # 找不到本地文件就跳过这行图片标记，不写乱码
                    continue
                items.append(("block", self._md_line_to_block(line)))
        return items

    def _resolve_image_path(self, url: str) -> str | None:
        """把 markdown 里的图片引用映射回本地文件路径"""
        name = url.split("/")[-1].split("?")[0]
        for cand in (OUTPUT_DIR / "images" / name, Path(url)):
            if cand.exists():
                return str(cand)
        return None

    def _parse_html_table_grid(self, html: str) -> dict:
        """解析 HTML 表格(含 rowspan/colspan)，返回网格布局:
        {n_rows, n_cols, cells:[{row,col,rs,cs,text}]}，按 HTML 表格规则计算每格位置。"""
        occupied = set()
        cells = []
        trs = re.findall(r"<tr>(.*?)</tr>", html, flags=re.DOTALL)
        max_col = 0
        for r, tr in enumerate(trs):
            c = 0
            for mtd in re.finditer(r"<t[dh]([^>]*)>(.*?)</t[dh]>", tr, flags=re.DOTALL):
                attrs, inner = mtd.group(1), mtd.group(2)
                def _attr(name):
                    mm = re.search(name + r'\s*=\s*"?(\d+)', attrs)
                    return max(1, int(mm.group(1))) if mm else 1
                rs, cs = _attr("rowspan"), _attr("colspan")
                while (r, c) in occupied:
                    c += 1
                text = re.sub(r"<[^>]+>", "", inner).strip()
                cells.append({"row": r, "col": c, "rs": rs, "cs": cs, "text": text})
                for dr in range(rs):
                    for dc in range(cs):
                        occupied.add((r + dr, c + dc))
                c += cs
                max_col = max(max_col, c)
        n_rows = len(trs)
        n_cols = max([cc for (_, cc) in occupied], default=-1) + 1
        n_cols = max(n_cols, max_col)
        return {"n_rows": n_rows, "n_cols": n_cols, "cells": cells}

    def _els(self, content: str) -> list:
        """文本 -> 飞书 elements，自动把 $..$ / $$..$$ 转成公式元素"""
        parts = re.split(r"(\$\$.+?\$\$|\$[^$\n]+?\$)", content)
        els = []
        for p in parts:
            if not p:
                continue
            if p.startswith("$$") and p.endswith("$$") and len(p) > 4:
                els.append({"equation": {"content": p[2:-2].strip()}})
            elif p.startswith("$") and p.endswith("$") and len(p) > 2:
                els.append({"equation": {"content": p[1:-1].strip()}})
            else:
                els.append({"text_run": {"content": p}})
        return els or [{"text_run": {"content": content}}]

    def _text_block(self, content: str) -> dict:
        return {"block_type": 2, "text": {"elements": self._els(content)}}

    def _md_line_to_block(self, line: str) -> dict:
        """将单行 Markdown 转为飞书 block（标题/分割线/引用/列表/公式/正文）"""
        # 分割线
        if re.fullmatch(r"(-{3,}|\*{3,}|_{3,})", line.strip()):
            return {"block_type": 22, "divider": {}}
        heading_match = re.match(r'^(#{1,6})\s+(.+)', line)
        if heading_match:
            level = min(len(heading_match.group(1)), 9)
            return {"block_type": level + 2,
                    f"heading{level}": {"elements": self._els(heading_match.group(2))}}
        # 引用：> 开头
        mq = re.match(r'^\s*[>＞]\s*(.+)', line)
        if mq:
            return {"block_type": 15, "quote": {"elements": self._els(mq.group(1))}}
        # 一级序列：1. / 2、 开头 → 有序列表(飞书自动编号，去掉原数字)
        m = re.match(r'^\s*(\d+)[.、．)]\s*(.+)', line)
        if m:
            return {"block_type": 13, "ordered": {"elements": self._els(m.group(2))}}
        # 二级序列：a. / b. 开头 → 子项(圆点，保留字母前缀以体现子序列)
        m = re.match(r'^\s*([a-zA-Z])[.、．)]\s*(.+)', line)
        if m:
            return {"block_type": 12, "bullet": {"elements": self._els(f"{m.group(1)}. {m.group(2)}")}}
        return self._text_block(line)


feishu_client = FeishuClient()

# ============================================================
# 网页文章解析 (公众号等)
# ============================================================
#
# 公众号文章是原生 HTML，文字/图片/表格都是真实结构，无需 OCR。
# 抓取 HTML -> 解析为有序条目(飞书块/表格/图片) -> 复用飞书写入逻辑。

WEB_HEADERS = {
    "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                  "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120 Safari/537.36",
    "Referer": "https://mp.weixin.qq.com/",
}

_BLOCK_TAGS = {"p", "h1", "h2", "h3", "h4", "h5", "h6", "li", "blockquote",
               "section", "div", "ul", "ol", "table", "img", "pre"}


# 飞书文字色板：1红 2橙 3黄 4绿 5蓝 6紫 7灰（参考 RGB 用于就近匹配）
_FEISHU_COLORS = {1: (216, 57, 49), 2: (222, 120, 2), 3: (220, 155, 4),
                  4: (47, 158, 68), 5: (24, 110, 232), 6: (123, 80, 198), 7: (140, 140, 140)}


def _parse_rgb(val: str):
    """把 CSS 颜色(hex / rgb()) 解析成 (r,g,b)，失败返回 None"""
    val = val.strip().lower()
    m = re.match(r"#([0-9a-f]{3})$", val)
    if m:
        h = m.group(1)
        return tuple(int(c * 2, 16) for c in h)
    m = re.match(r"#([0-9a-f]{6})$", val)
    if m:
        h = m.group(1)
        return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))
    m = re.match(r"rgba?\(([^)]+)\)", val)
    if m:
        parts = [p.strip() for p in m.group(1).split(",")]
        try:
            return tuple(int(float(parts[i])) for i in range(3))
        except Exception:
            return None
    return None


def _nearest_feishu_color(rgb, is_bg=False):
    """RGB 映射到飞书色枚举(按色相,浅色也能认)。黑/白/无色不上色返回 None。"""
    if rgb is None:
        return None
    import colorsys
    r, g, b = rgb
    mx, mn = max(rgb), min(rgb)
    if not is_bg and mx < 80:      # 接近黑=正文默认色，不上色
        return None
    if is_bg and mn > 240:         # 接近白=无高亮
        return None
    if mx - mn < 18:               # 几乎无彩(灰)
        if not is_bg and mn > 220: # 接近白的文字，不上色
            return None
        return 7                   # 灰
    deg = colorsys.rgb_to_hsv(r/255, g/255, b/255)[0] * 360
    if deg < 20 or deg >= 330: return 1   # 红/粉
    if deg < 45:  return 2                 # 橙
    if deg < 70:  return 3                 # 黄
    if deg < 170: return 4                 # 绿
    if deg < 255: return 5                 # 蓝
    if deg < 290: return 6                 # 紫
    return 1                               # 品红→红


def _style_colors(node):
    """从元素 style 属性解析出 (文字色枚举, 背景色枚举)"""
    style = node.get("style", "") if hasattr(node, "get") else ""
    if not style:
        return None, None
    tc = bg = None
    mc = re.search(r"(?<!-)color\s*:\s*([^;]+)", style)
    if mc:
        tc = _nearest_feishu_color(_parse_rgb(mc.group(1)), is_bg=False)
    mb = re.search(r"background(?:-color)?\s*:\s*([^;]+)", style)
    if mb:
        bg = _nearest_feishu_color(_parse_rgb(mb.group(1)), is_bg=True)
    return tc, bg


def _inline_elements(el) -> list[dict]:
    """把块元素的内联内容转为飞书 text elements，保留：加粗/斜体/下划线/删除线/
    内联代码/文字颜色/高亮/超链接。"""
    from bs4 import NavigableString
    from urllib.parse import quote
    elements: list[dict] = []

    def rec(node, st):
        for c in node.children:
            if isinstance(c, NavigableString):
                t = str(c).replace("\xa0", " ")
                if t.strip():
                    style = {k: v for k, v in st.items() if v and k != "link"}
                    if st.get("link"):
                        style["link"] = {"url": quote(st["link"], safe="")}
                    run = {"content": t}
                    if style:
                        run["text_element_style"] = style
                    elements.append({"text_run": run})
            elif getattr(c, "name", None):
                ns = dict(st)
                n = c.name
                if n in ("strong", "b"): ns["bold"] = True
                if n in ("em", "i"): ns["italic"] = True
                if n in ("u", "ins"): ns["underline"] = True
                if n in ("s", "del", "strike"): ns["strikethrough"] = True
                if n == "code": ns["inline_code"] = True
                if n == "a" and c.get("href", "").startswith("http"): ns["link"] = c.get("href")
                tc, cbg = _style_colors(c)
                if tc: ns["text_color"] = tc
                if cbg: ns["background_color"] = cbg
                rec(c, ns)

    rec(el, {})
    if not elements:
        txt = el.get_text(" ", strip=True)
        if txt:
            elements = [{"text_run": {"content": txt}}]
    return elements


async def _download_image(client, url: str) -> str | None:
    """下载远程图片到本地临时文件，返回路径"""
    try:
        r = await client.get(url, headers=WEB_HEADERS, follow_redirects=True, timeout=30)
        if r.status_code != 200 or not r.content:
            return None
        ext = ".jpg"
        u = url.lower()
        if "wx_fmt=png" in u or ".png" in u:
            ext = ".png"
        elif "wx_fmt=gif" in u or ".gif" in u:
            ext = ".gif"
        path = UPLOAD_DIR / f"web_{uuid.uuid4().hex[:10]}{ext}"
        path.write_bytes(r.content)
        return str(path)
    except Exception as e:
        logger.warning(f"下载图片失败 {url[:60]}: {e}")
        return None


def _check_public_url(url: str):
    """只允许 http/https 的公网地址，阻断内网/本机(防 SSRF 探测内网)。"""
    from urllib.parse import urlparse
    import ipaddress, socket
    u = urlparse(url)
    if u.scheme not in ("http", "https"):
        raise RuntimeError("只支持 http/https 链接")
    host = u.hostname or ""
    try:
        ip = ipaddress.ip_address(socket.gethostbyname(host))
        if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_reserved:
            raise RuntimeError("不允许访问内网/本机地址")
    except (socket.gaierror, ValueError):
        raise RuntimeError("无法解析该链接的地址")


async def parse_web_article(url: str) -> tuple[str, list[tuple]]:
    """抓取网页文章，返回 (标题, 有序条目列表)。条目格式与飞书写入一致。
    抓到空/被反爬时自动重试几次。"""
    from bs4 import BeautifulSoup
    _check_public_url(url)

    async with httpx.AsyncClient(timeout=40) as client:
        content = soup = None
        last_err = ""
        for attempt in range(3):  # 失败/被反爬自动重试
            try:
                resp = await client.get(url, headers=WEB_HEADERS, follow_redirects=True)
                if resp.status_code == 200:
                    soup = BeautifulSoup(resp.text, "html.parser")
                    content = soup.find(id="js_content") or soup.find("article") or soup.body
                    if content and content.get_text(strip=True):
                        break
                    last_err = "抓到空内容(可能被反爬)"
                else:
                    last_err = f"HTTP {resp.status_code}"
            except Exception as e:
                last_err = str(e)
            await asyncio.sleep(1.5 * (attempt + 1))  # 递增等待后重试
        if content is None or not content.get_text(strip=True):
            raise RuntimeError(f"抓取文章失败: {last_err}")

        title_el = soup.find(id="activity-name") or soup.find("h1") or soup.find("title")
        title = title_el.get_text(strip=True) if title_el else "网页文章"

        items: list[tuple] = []
        seen_imgs: set[str] = set()

        async def emit_image(img):
            src = img.get("data-src") or img.get("src") or ""
            if not src or src in seen_imgs or src.startswith("data:"):
                return
            seen_imgs.add(src)
            local = await _download_image(client, src)
            if local:
                items.append(("image", local))

        def heading_item(el):
            level = min(int(el.name[1]), 9)
            return ("block", {"block_type": level + 2,
                              f"heading{level}": {"elements": _inline_elements(el)}})

        def text_item(el, block_type=2, key="text"):
            els = _inline_elements(el)
            if not els:
                return None
            return ("block", {"block_type": block_type, key: {"elements": els}})

        def has_block_child(el):
            return el.find(_BLOCK_TAGS - {"img"}) is not None

        async def walk(node):
            for child in node.children:
                name = getattr(child, "name", None)
                if not name:
                    continue
                if name == "img":
                    await emit_image(child)
                elif name == "hr":
                    items.append(("block", {"block_type": 22, "divider": {}}))
                elif name == "blockquote":
                    els = _inline_elements(child)
                    if els:
                        items.append(("block", {"block_type": 15, "quote": {"elements": els}}))
                    for im in child.find_all("img"):
                        await emit_image(im)
                elif name in ("h1", "h2", "h3", "h4", "h5", "h6"):
                    items.append(heading_item(child))
                elif name == "table":
                    rows = []
                    for tr in child.find_all("tr"):
                        cells = [c.get_text(" ", strip=True) for c in tr.find_all(["td", "th"])]
                        if cells:
                            rows.append(cells)
                    if rows:
                        items.append(("table", rows))
                elif name in ("ul", "ol"):
                    # 有序列表(ol)用飞书有序列表块(13,自动编号)，无序(ul)用圆点(12)
                    bt, key = (13, "ordered") if name == "ol" else (12, "bullet")
                    for li in child.find_all("li", recursive=False):
                        it = text_item(li, bt, key)
                        if it:
                            items.append(it)
                        for im in li.find_all("img"):
                            await emit_image(im)
                elif name == "p":
                    it = text_item(child)
                    if it:
                        items.append(it)
                    for im in child.find_all("img"):
                        await emit_image(im)
                elif name == "pre":
                    # 代码块：每个 <code> 子元素是一行，保留换行
                    codes = child.find_all("code", recursive=False)
                    if codes:
                        code_text = "\n".join(c.get_text() for c in codes)
                    else:
                        code_text = child.get_text()
                    code_text = code_text.replace("\xa0", " ").rstrip()
                    if code_text:
                        items.append(("block", {
                            "block_type": 14,
                            "code": {"elements": [{"text_run": {"content": code_text}}],
                                     "style": {"language": 1, "wrap": True}},
                        }))
                else:  # section/div/span 等容器
                    if has_block_child(child):
                        await walk(child)
                    else:
                        it = text_item(child)
                        if it:
                            items.append(it)
                        for im in child.find_all("img"):
                            await emit_image(im)

        await walk(content)

        # 过滤掉空文本/空标题块
        def is_empty_block(item):
            kind, payload = item
            if kind != "block":
                return False
            body = list(payload.values())[1]
            els = body.get("elements", [])
            return not any(e.get("text_run", {}).get("content", "").strip() for e in els)

        items = [it for it in items if not is_empty_block(it)]
        return title, items


# ============================================================
# OCR 引擎 (MinerU)
# ============================================================
#
# 说明：旧版用 PaddleOCR + PPStructure，但版面分析在长截图/非标准排版上
# 经常漏检整段文字、且会把表格拆成碎片。改用 MinerU(pipeline 后端)，它专门
# 做「文档图片/PDF → 带表格结构的 Markdown」，本地离线运行、表格还原准确。
#
# 调用方式：通过 mineru CLI 子进程处理，输出目录里取 <name>.md。
# 注意：mineru 内部会起一个本机 FastAPI 子服务，若系统设了 http_proxy 会把
# 本机请求也走代理导致超时，因此子进程里强制 NO_PROXY=127.0.0.1,localhost。

def _heading_level_from_text(text: str):
    """根据标题里的编号判断层级：一二三/第N章/单数字→1，N.N→2，N.N.N→3；判断不了返回 None"""
    t = text.strip()
    m = re.match(r"^(\d+(?:\.\d+)+)", t)        # 形如 2.1 / 3.2.1（多级编号）
    if m:
        return min(m.group(1).count(".") + 1, 6)  # 2.1→2, 3.2.1→3
    if re.match(r"^[一二三四五六七八九十百]+[、，,.．)）]", t):
        return 1
    if re.match(r"^第[一二三四五六七八九十\d]+[章节部分讲篇]", t):
        return 1
    if re.match(r"^\d+[、.．)）]", t):           # 单个数字开头 1. / 2、 → 顶层
        return 1
    return None


def _looks_like_english_title(text: str) -> bool:
    """判断是否英文为主的标题(论文名常见)，用于识别被误判成标题的列表项"""
    cjk = sum(1 for ch in text if "一" <= ch <= "鿿")
    letters = sum(1 for ch in text if ch.isascii() and ch.isalpha())
    return letters >= 4 and cjk <= 1


def normalize_heading_levels(markdown: str) -> str:
    """修正 MinerU 的标题误判：
    - 真章节(中文编号标题)按编号定级(一二三/N → 1, N.N → 2 ...)
    - 被误判成标题的「英文数字标题」(论文)降级为普通行 → 后续转有序列表
    - 被误判成标题的「字母a./b.开头」(子项)降级为普通行 → 后续转子项"""
    out = []
    for line in markdown.split("\n"):
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if not m:
            out.append(line); continue
        text = m.group(2).strip()
        # 字母开头(a. b.)被误判为标题 -> 降为普通行
        if re.match(r"^[a-zA-Z][.、．)]", text):
            out.append(text); continue
        # 数字+英文标题(论文)被误判为标题 -> 降为普通行
        mm = re.match(r"^\d+[.、．)]\s*(.+)", text)
        if mm and _looks_like_english_title(mm.group(1)):
            out.append(text); continue
        # 真标题：按编号定级
        lvl = _heading_level_from_text(text)
        out.append(("#" * lvl + " " + text) if lvl else line)
    return "\n".join(out)


class OCREngine:
    def __init__(self):
        self._initialized = False
        self._mineru_cli = None

    def initialize(self):
        if self._initialized:
            return
        cli = shutil.which("mineru")
        if not cli:
            # 退而求其次：用当前 Python 解释器同目录下的 mineru
            cand = Path(sys.executable).parent / "mineru"
            if cand.exists():
                cli = str(cand)
        if not cli:
            raise RuntimeError("MinerU 未安装，请运行: pip install -U 'mineru[all]'")
        self._mineru_cli = cli
        self._initialized = True
        logger.info(f"✅ OCR 引擎就绪 (MinerU: {cli})")

    def process_image(self, image_path: str, dewatermark: bool = False) -> dict:
        self.initialize()
        start_time = time.time()
        image_path = Path(image_path)

        # 去水印(浅色水印)：把接近白色的像素刷成纯白，水印消失而黑字/彩色不变
        if dewatermark:
            cleaned = self._dewatermark(image_path)
            if cleaned:
                image_path = cleaned

        # 长图先在空白行处切成多段，逐段识别再拼接，避免缩放导致漏字
        segments, tmp_segments = self._split_tall_image(image_path)
        try:
            md_parts = []
            for seg in segments:
                md_parts.append(self._run_mineru(seg))
            markdown = "\n\n".join(p for p in md_parts if p.strip())
            markdown = normalize_heading_levels(markdown)  # 按编号还原标题层级
            elapsed = time.time() - start_time
            return {
                "markdown": markdown,
                "regions": [],
                "stats": {
                    "elapsed": round(elapsed, 2),
                    "region_count": markdown.count("\n\n") + 1,
                    "table_count": markdown.count("<table"),
                    "image_count": markdown.count("!["),
                    "segments": len(segments),
                },
            }
        finally:
            for p in tmp_segments:
                try: os.remove(p)
                except OSError: pass

    def _run_mineru(self, image_path) -> str:
        """对单张图片跑 MinerU，返回 markdown"""
        out_dir = OUTPUT_DIR / f".mineru_{uuid.uuid4().hex[:8]}"
        out_dir.mkdir(parents=True, exist_ok=True)

        env = os.environ.copy()
        env["NO_PROXY"] = "127.0.0.1,localhost"
        env["no_proxy"] = "127.0.0.1,localhost"
        # 离线模式：模型已预下载，强制只用本地缓存，断网也能识别(且更快)
        env["HF_HUB_OFFLINE"] = "1"
        env["TRANSFORMERS_OFFLINE"] = "1"
        env["HF_DATASETS_OFFLINE"] = "1"

        try:
            proc = subprocess.run(
                [self._mineru_cli, "-p", str(image_path), "-o", str(out_dir),
                 "-b", "pipeline", "-l", "ch"],
                capture_output=True, text=True, env=env, timeout=600,
            )
            if proc.returncode != 0:
                logger.error(f"MinerU 处理失败: {proc.stderr[-500:]}")
                raise RuntimeError(f"MinerU 处理失败 (exit {proc.returncode})")

            md_files = list(out_dir.glob("**/*.md"))
            if not md_files:
                raise RuntimeError("MinerU 未生成 Markdown 结果")
            md_path = md_files[0]
            markdown = md_path.read_text(encoding="utf-8")
            return self._relocate_images(markdown, md_path.parent)
        finally:
            shutil.rmtree(out_dir, ignore_errors=True)

    def _dewatermark(self, image_path: Path, thresh: int = 225) -> Path | None:
        """去浅色水印：三通道都 >= thresh 的像素刷成纯白(水印盖在白底上的典型情况)。
        黑字/彩色内容低于阈值，不受影响。返回处理后的临时图路径。"""
        try:
            from PIL import Image
            import numpy as np
            img = Image.open(image_path).convert("RGB")
            a = np.array(img)
            mask = (a[:, :, 0] >= thresh) & (a[:, :, 1] >= thresh) & (a[:, :, 2] >= thresh)
            a[mask] = [255, 255, 255]
            out = UPLOAD_DIR / f"dewm_{uuid.uuid4().hex[:8]}.png"
            Image.fromarray(a).save(out)
            logger.info(f"已去水印(刷白 {round(mask.mean()*100,1)}% 近白像素)")
            return out
        except Exception as e:
            logger.warning(f"去水印失败，用原图: {e}")
            return None

    def _split_tall_image(self, image_path: Path):
        """长图按空白行切段。返回 (段路径列表, 需清理的临时文件列表)。
        非长图直接返回原图，不切。"""
        try:
            from PIL import Image
            import numpy as np
        except ImportError:
            return [str(image_path)], []

        img = Image.open(image_path).convert("RGB")
        w, h = img.size
        # 仅对"又长又窄"的图分段(高>3000 且 高宽比>2.5)
        if h <= 3000 or h / max(w, 1) < 2.5:
            return [str(image_path)], []

        gray = np.asarray(img.convert("L"))
        row_brightness = gray.mean(axis=1)  # 每行平均亮度，越大越接近空白

        seg_target = 2200      # 每段目标高度
        search = 400           # 在目标边界±search内找最空白的行作为切点
        cuts = [0]
        y = 0
        while h - y > seg_target + search:
            lo = y + seg_target - search
            hi = y + seg_target + search
            cut = lo + int(np.argmax(row_brightness[lo:hi]))
            cuts.append(cut)
            y = cut
        cuts.append(h)

        seg_paths, tmp = [], []
        for i in range(len(cuts) - 1):
            top, bot = cuts[i], cuts[i + 1]
            crop = img.crop((0, top, w, bot))
            p = str(UPLOAD_DIR / f"seg_{uuid.uuid4().hex[:8]}.png")
            crop.save(p)
            seg_paths.append(p)
            tmp.append(p)
        logger.info(f"长图({w}x{h})切成 {len(seg_paths)} 段识别")
        return seg_paths, tmp

    def _relocate_images(self, markdown: str, md_dir: Path) -> str:
        """把 MinerU 抽取的图片移到 outputs/images 下，并改写 markdown 里的引用路径。
        保留所有图片(不过滤)，并自动裁掉四周的纯色空白边(避免飞书里图片下方一大块空白)。"""
        src_img_dir = md_dir / "images"
        if not src_img_dir.exists():
            return markdown
        dst_img_dir = OUTPUT_DIR / "images"
        dst_img_dir.mkdir(parents=True, exist_ok=True)
        for img in src_img_dir.iterdir():
            dst = dst_img_dir / img.name
            try:
                if not self._trim_to(img, dst):   # 裁白边并保存到 dst
                    shutil.copy(img, dst)          # 裁剪失败就原样拷
                markdown = markdown.replace(f"images/{img.name}", f"/api/image/{img.name}")
            except OSError:
                pass
        return markdown

    def _trim_to(self, src: Path, dst: Path) -> bool:
        """裁掉图片四周的纯色(近白)空白边，保存到 dst。成功返回 True。"""
        try:
            from PIL import Image
            import numpy as np
            im = Image.open(src).convert("RGB")
            a = np.asarray(im.convert("L"))
            mask = a < 245  # 非近白像素=内容
            if not mask.any():
                return False
            rows = np.where(mask.any(axis=1))[0]
            cols = np.where(mask.any(axis=0))[0]
            top, bot = int(rows[0]), int(rows[-1]) + 1
            left, right = int(cols[0]), int(cols[-1]) + 1
            # 留 4px 边距，避免裁得太紧
            m = 4
            top, left = max(0, top - m), max(0, left - m)
            bot, right = min(im.size[1], bot + m), min(im.size[0], right + m)
            im.crop((left, top, right, bot)).save(dst)
            return True
        except Exception:
            return False


ocr_engine = OCREngine()

# ============================================================
# 导出
# ============================================================

def export_to_docx(markdown_text: str, output_path: str) -> bool:
    """用 python-docx 生成 Word，纯本地、不依赖 pandoc。支持标题/表格/列表/图片。"""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        logger.error("python-docx 未安装，请运行: pip install python-docx")
        return False

    try:
        doc = Document()
        # 按 <table> 切分，逐段处理
        parts = re.split(r"(<table>.*?</table>)", markdown_text, flags=re.DOTALL)
        for part in parts:
            if not part.strip():
                continue
            if part.strip().startswith("<table>"):
                rows = []
                for tr in re.findall(r"<tr>(.*?)</tr>", part, flags=re.DOTALL):
                    cells = re.findall(r"<t[dh][^>]*>(.*?)</t[dh]>", tr, flags=re.DOTALL)
                    cells = [re.sub(r"<[^>]+>", "", c).strip() for c in cells]
                    if cells:
                        rows.append(cells)
                if rows:
                    ncol = max(len(r) for r in rows)
                    t = doc.add_table(rows=0, cols=ncol)
                    t.style = "Table Grid"
                    for r in rows:
                        cells = t.add_row().cells
                        for i in range(ncol):
                            cells[i].text = r[i] if i < len(r) else ""
                continue
            for line in part.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if re.fullmatch(r"<!--.*-->", line) or re.fullmatch(r"-{3,}", line):
                    continue
                # 图片：按原始分辨率放，最多 5.5 英寸宽，绝不放大(否则小图会糊/变形)
                img = re.search(r"!\[[^\]]*\]\(([^)]+)\)", line)
                if img:
                    name = img.group(1).split("/")[-1].split("?")[0]
                    p = OUTPUT_DIR / "images" / name
                    if p.exists():
                        try:
                            from docx.shared import Inches
                            from PIL import Image as _PILImage
                            with _PILImage.open(p) as _im:
                                px_w = _im.size[0]
                            width_in = min(px_w / 96.0, 5.5)  # 96dpi 换算，封顶5.5寸
                            doc.add_picture(str(p), width=Inches(width_in))
                        except Exception:
                            try:
                                doc.add_picture(str(p))  # 兜底:原尺寸
                            except Exception:
                                pass
                    continue
                # 标题
                h = re.match(r"^(#{1,6})\s+(.+)", line)
                if h:
                    doc.add_heading(h.group(2), level=min(len(h.group(1)), 4))
                else:
                    doc.add_paragraph(line)
        doc.save(output_path)
        return True
    except Exception as e:
        logger.error(f"生成 Word 失败: {e}", exc_info=True)
        return False

# ============================================================
# FastAPI
# ============================================================

# 识别结果/缓存保留天数，超过自动清理(隐私+磁盘)。可用环境变量覆盖。
CLEANUP_DAYS = int(os.environ.get("DOCLENS_KEEP_DAYS", "7"))


def cleanup_old_files(days: int = CLEANUP_DAYS):
    """清理超过 N 天的上传/输出/抠图文件。"""
    if days <= 0:
        return
    cutoff = time.time() - days * 86400
    removed = 0
    for base in (UPLOAD_DIR, OUTPUT_DIR, OUTPUT_DIR / "images"):
        if not base.exists():
            continue
        for p in base.iterdir():
            if p.is_file():
                try:
                    if p.stat().st_mtime < cutoff:
                        p.unlink(); removed += 1
                except OSError:
                    pass
    if removed:
        logger.info(f"已清理 {removed} 个过期文件(>{days}天)")


@asynccontextmanager
async def lifespan(app: FastAPI):
    status = network_guard.get_status()
    if status["online"]:
        logger.warning("⚠️  网络已连接 — 本地模式下建议断网")
    else:
        logger.info("✅ 离线模式运行中")

    cleanup_old_files()  # 启动时清理过期文件

    # 加载飞书配置
    cfg = load_config()
    if cfg.get("feishu_app_id"):
        feishu_client.configure(cfg["feishu_app_id"], cfg["feishu_app_secret"])
        logger.info("飞书配置已加载")

    logger.info(f"🚀 DocLens 启动: http://{HOST}:{PORT}")
    yield
    logger.info("DocLens 已停止")


app = FastAPI(title="DocLens", lifespan=lifespan)
tasks: dict = {}

_ALLOWED_ORIGINS = {f"http://{HOST}:{PORT}", "http://localhost:" + str(PORT)}


@app.middleware("http")
async def _csrf_guard(request, call_next):
    """阻止你访问的其他网站偷偷调用本机接口(CSRF)：
    凡是带外部 Origin 的写操作(POST/PUT/PATCH/DELETE)一律拒绝。
    本应用自己的页面同源(Origin 匹配)，正常放行。"""
    if request.method not in ("GET", "HEAD", "OPTIONS"):
        origin = request.headers.get("origin")
        if origin and origin not in _ALLOWED_ORIGINS:
            return JSONResponse({"detail": "跨站请求被拒绝"}, status_code=403)
    return await call_next(request)


# ---- 页面 ----

@app.get("/", response_class=HTMLResponse)
async def serve_index():
    index_path = STATIC_DIR / "index.html"
    if not index_path.exists():
        raise HTTPException(500, "前端文件缺失")
    return HTMLResponse(index_path.read_text(encoding="utf-8"))


# ---- 网络状态 ----

@app.get("/api/network")
async def get_network_status():
    return network_guard.get_status()


# ---- 飞书配置 ----

@app.get("/api/feishu/config")
async def get_feishu_config():
    return {
        "configured": feishu_client.is_configured,
        "app_id_masked": feishu_client.app_id[:4] + "****" if feishu_client.app_id else "",
    }


@app.post("/api/feishu/config")
async def set_feishu_config(body: dict = Body(...)):
    app_id = body.get("app_id", "").strip()
    app_secret = body.get("app_secret", "").strip()
    if not app_id or not app_secret:
        raise HTTPException(400, "App ID 和 App Secret 不能为空")

    cfg = load_config()
    cfg["feishu_app_id"] = app_id
    cfg["feishu_app_secret"] = app_secret
    save_config(cfg)
    feishu_client.configure(app_id, app_secret)

    # 验证凭证
    try:
        await feishu_client._ensure_token()
        return {"status": "ok", "message": "飞书配置验证成功"}
    except Exception as e:
        return {"status": "error", "message": f"验证失败: {str(e)}"}


# ---- 飞书账号登录(OAuth)：让生成的文档归你所有 ----

def _redirect_uri(request) -> str:
    return f"http://{HOST}:{PORT}/api/auth/callback"


_oauth_state = ""


@app.get("/api/auth/login")
async def auth_login():
    global _oauth_state
    if not feishu_client.is_configured:
        raise HTTPException(400, "请先配置飞书应用凭证")
    _oauth_state = uuid.uuid4().hex  # 防 CSRF：随机 state，回调时校验
    return RedirectResponse(feishu_client.authorize_url(_redirect_uri(None), _oauth_state))


@app.get("/api/auth/callback")
async def auth_callback(code: str = "", state: str = "", error: str = ""):
    if error or not code:
        return HTMLResponse(f"<script>alert('登录失败: {error or '未收到授权码'}');location.href='/'</script>")
    if not _oauth_state or state != _oauth_state:
        return HTMLResponse("<script>alert('登录校验失败(state不匹配)，请重新点登录');location.href='/'</script>")
    try:
        await feishu_client.exchange_code(code, _redirect_uri(None))
        name = feishu_client.user_name or "你"
        return HTMLResponse(f"<script>alert('✅ 已登录飞书：{name}，以后生成的文档都归你所有');location.href='/'</script>")
    except Exception as e:
        return HTMLResponse(f"<script>alert('登录失败: {str(e)}');location.href='/'</script>")


@app.get("/api/auth/status")
async def auth_status():
    return {"logged_in": feishu_client.has_user_auth, "name": feishu_client.user_name}


@app.post("/api/auth/logout")
async def auth_logout():
    feishu_client.logout()
    return {"status": "ok"}


# ---- 本地上传转换 ----

@app.post("/api/convert/local")
async def convert_local(
    file: list[UploadFile] = File(...),
    output_format: str = Form("md"),
    watermark: str = Form(""),
    dewatermark: str = Form(""),
):
    allowed = {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif"}
    files = file if isinstance(file, list) else [file]
    task_id = str(uuid.uuid4())[:8]

    upload_paths = []
    for idx, f in enumerate(files):
        ext = Path(f.filename).suffix.lower()
        if ext not in allowed:
            raise HTTPException(400, f"不支持的格式: {ext}")
        p = UPLOAD_DIR / f"{task_id}_{idx}_{f.filename}"
        with open(p, "wb") as out:
            out.write(await f.read())
        upload_paths.append(str(p))

    tasks[task_id] = {
        "id": task_id,
        "mode": "local",
        "filename": files[0].filename + (f" 等{len(files)}张" if len(files) > 1 else ""),
        "status": "processing",
        "progress": "准备中...",
        "created_at": datetime.now().isoformat(),
        "result": None,
    }

    asyncio.get_event_loop().run_in_executor(
        None, _process_local_task, task_id, upload_paths, output_format, watermark,
        dewatermark in ("1", "true", "on", "yes")
    )
    return {"task_id": task_id, "status": "processing"}


def strip_watermark(markdown: str, watermark: str) -> str:
    """安全去水印词：只删「独立出现」(前后是空白/边界)的水印片段，
    不动句子里嵌着的同款词，避免误删正文。watermark 可逗号分隔多个。
    OCR 水印通常是散落的孤立片段，所以这样既能去水印又最大限度保护正文。"""
    if not watermark.strip():
        return markdown
    words = {w.strip() for w in re.split(r"[,，]", watermark) if w.strip()}
    if not words:
        return markdown
    out_lines = []
    for line in markdown.split("\n"):
        # 按空白切分(保留分隔符)，只丢掉「整段等于水印词」的孤立 token
        tokens = re.split(r"(\s+)", line)
        kept = [t for t in tokens if t.strip() not in words]
        out_lines.append("".join(kept))
    md = "\n".join(out_lines)
    md = re.sub(r"[ \t]{2,}", " ", md)
    md = re.sub(r" +\n", "\n", md)
    return md


def _process_local_task(task_id: str, image_paths, output_format: str, watermark: str = "", dewatermark: bool = False):
    # image_paths 可能是单个路径(str)或多个路径(list)；多张=同一篇文章的多段，按顺序拼接
    if isinstance(image_paths, str):
        image_paths = [image_paths]
    try:
        md_parts, total_regions, total_tables, total_images = [], 0, 0, 0
        for i, img in enumerate(image_paths):
            if len(image_paths) > 1:
                tasks[task_id]["progress"] = f"正在识别第 {i+1}/{len(image_paths)} 张..."
            else:
                tasks[task_id]["progress"] = "正在识别文档内容..."
            r = ocr_engine.process_image(img, dewatermark=dewatermark)
            md_parts.append(r["markdown"])
            total_regions += r["stats"]["region_count"]
            total_tables += r["stats"]["table_count"]
            total_images += r["stats"]["image_count"]
        # 多段直接顺序拼接(同一篇文章，不加分隔标记)
        combined = "\n\n".join(p for p in md_parts if p.strip())
        result = {"markdown": strip_watermark(combined, watermark),
                  "stats": {"elapsed": 0, "region_count": total_regions,
                            "table_count": total_tables, "image_count": total_images,
                            "segments": len(image_paths)}}

        md_filename = f"{task_id}_result.md"
        md_path = OUTPUT_DIR / md_filename
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(result["markdown"])

        tasks[task_id]["result"] = {
            "markdown": result["markdown"],
            "stats": result["stats"],
            "md_file": md_filename,
            "docx_file": None,
            "feishu_url": None,
        }

        if output_format == "docx":
            tasks[task_id]["progress"] = "正在生成 Word 文档..."
            docx_filename = f"{task_id}_result.docx"
            docx_path = OUTPUT_DIR / docx_filename
            if export_to_docx(result["markdown"], str(docx_path)):
                tasks[task_id]["result"]["docx_file"] = docx_filename

        elif output_format == "feishu":
            # 源文件已在本机识别完成，这里只把识别结果(文字/表格/抠出的图)写入飞书
            if not feishu_client.is_configured:
                raise RuntimeError("未配置飞书凭证，无法上传到飞书")
            if not feishu_client.has_user_auth:
                raise RuntimeError("请先登录飞书账号，以确保文档归你所有")
            tasks[task_id]["progress"] = "正在写入飞书文档..."
            title = tasks[task_id].get("filename") or f"DocLens 识别结果 {datetime.now().strftime('%m-%d %H:%M')}"

            async def _do():
                url = await feishu_client.create_doc_with_content(title, result["markdown"])
                await feishu_client.append_index_record(title, url, "本地上传")
                return url

            feishu_url = asyncio.run(_do())
            tasks[task_id]["result"]["feishu_url"] = feishu_url

        tasks[task_id]["status"] = "completed"
        tasks[task_id]["progress"] = "完成"
    except Exception as e:
        logger.error(f"任务 {task_id} 失败: {e}", exc_info=True)
        tasks[task_id]["status"] = "failed"
        tasks[task_id]["progress"] = f"失败: {str(e)}"
    finally:
        for p in image_paths:
            try: os.remove(p)
            except OSError: pass


# ---- 批量转换(按文件夹分组：散图=各自一篇，子文件夹=合并一篇) ----

def _group_images(saved):
    """saved: [(相对路径, 本地路径, 原文件名)] -> [(标题, [本地路径按序])]"""
    groups, order = {}, []
    for rel, path, fname in saved:
        segs = [s for s in rel.replace("\\", "/").split("/") if s]
        sub = segs[1:] if len(segs) >= 2 else segs   # 去掉最外层所选文件夹名
        if len(sub) >= 2:                            # 在子文件夹里 -> 合并成一篇
            key, sortk = sub[0], "/".join(sub[1:])
            title = sub[0]
        else:                                        # 直接散放 -> 各自一篇
            base = sub[0] if sub else fname
            key, sortk, title = "__one__/" + base, base, Path(base).stem
        if key not in groups:
            groups[key] = {"title": title, "items": []}; order.append(key)
        groups[key]["items"].append((sortk, path))
    out = []
    for key in order:
        paths = [p for _, p in sorted(groups[key]["items"], key=lambda x: x[0])]
        out.append((groups[key]["title"], paths))
    return out


@app.post("/api/convert/batch")
async def convert_batch(
    file: list[UploadFile] = File(...),
    rel_paths: str = Form("[]"),
    output_format: str = Form("md"),
    watermark: str = Form(""),
    dewatermark: str = Form(""),
):
    allowed = {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif"}
    files = file if isinstance(file, list) else [file]
    try:
        rels = json.loads(rel_paths)
    except Exception:
        rels = []
    if len(rels) != len(files):
        rels = [f.filename for f in files]

    batch_id = str(uuid.uuid4())[:8]
    saved = []
    for idx, (f, rel) in enumerate(zip(files, rels)):
        if Path(f.filename).suffix.lower() not in allowed:
            continue  # 跳过非图片
        p = UPLOAD_DIR / f"{batch_id}_{idx}_{Path(f.filename).name}"
        with open(p, "wb") as o:
            o.write(await f.read())
        saved.append((rel, str(p), f.filename))

    if not saved:
        raise HTTPException(400, "没有可处理的图片")

    groups = _group_images(saved)
    tasks[batch_id] = {
        "id": batch_id, "mode": "batch", "status": "processing",
        "progress": f"共 {len(groups)} 篇，准备中...",
        "created_at": datetime.now().isoformat(),
        "result": {"items": [{"title": t, "status": "待处理"} for t, _ in groups],
                   "done": 0, "total": len(groups)},
    }
    asyncio.get_event_loop().run_in_executor(
        None, _process_batch_task, batch_id, groups, output_format, watermark,
        dewatermark in ("1", "true", "on", "yes")
    )
    return {"task_id": batch_id, "status": "processing", "total": len(groups)}


def _process_batch_task(batch_id, groups, output_format, watermark, dewatermark):
    items = tasks[batch_id]["result"]["items"]
    done = 0
    for i, (title, paths) in enumerate(groups):
        items[i]["status"] = "处理中"
        tasks[batch_id]["progress"] = f"第 {i+1}/{len(groups)} 篇：{title}"
        sub_id = f"{batch_id}s{i}"
        tasks[sub_id] = {"id": sub_id, "mode": "local", "filename": title,
                         "status": "processing", "progress": "", "result": None}
        _process_local_task(sub_id, list(paths), output_format, watermark, dewatermark)
        st = tasks.get(sub_id, {})
        if st.get("status") == "completed":
            r = st["result"]
            items[i].update(status="完成", feishu_url=r.get("feishu_url"),
                            md_file=r.get("md_file"), docx_file=r.get("docx_file"))
        else:
            items[i].update(status="失败", error=st.get("progress", "未知错误"))
        tasks.pop(sub_id, None)
        done += 1
        tasks[batch_id]["result"]["done"] = done
    ok = sum(1 for it in items if it["status"] == "完成")
    tasks[batch_id]["status"] = "completed"
    tasks[batch_id]["progress"] = f"完成：成功 {ok}/{len(groups)} 篇"


# ---- 飞书链接转换 ----

@app.post("/api/convert/feishu")
async def convert_feishu(body: dict = Body(...)):
    url = body.get("url", "").strip()
    if not url:
        raise HTTPException(400, "请提供飞书文档链接")

    if not feishu_client.is_configured:
        raise HTTPException(400, "请先配置飞书应用凭证")
    if not feishu_client.has_user_auth:
        raise HTTPException(400, "请先登录飞书账号(右上角『登录飞书』)，以确保生成的文档归你所有")

    if not network_guard.check_feishu_reachable():
        raise HTTPException(400, "连不上飞书服务器，请检查网络")

    task_id = str(uuid.uuid4())[:8]
    tasks[task_id] = {
        "id": task_id,
        "mode": "feishu",
        "filename": url,
        "status": "processing",
        "progress": "解析飞书链接...",
        "created_at": datetime.now().isoformat(),
        "result": None,
    }

    # 飞书任务需要异步执行
    asyncio.create_task(_process_feishu_task(task_id, url))
    return {"task_id": task_id, "status": "processing"}


async def _process_feishu_task(task_id: str, url: str):
    local_paths = []
    try:
        # 1. 解析链接
        tasks[task_id]["progress"] = "解析飞书文档链接..."
        doc_token, doc_type = feishu_client.parse_doc_url(url)

        # 2. 提取图片
        tasks[task_id]["progress"] = "从飞书文档下载图片..."
        local_paths = await feishu_client.extract_images_from_doc(doc_token)
        logger.info(f"从飞书文档提取了 {len(local_paths)} 张图片")

        # 3. 逐张 OCR (在线程中执行，避免阻塞)
        all_markdown = []
        for idx, img_path in enumerate(local_paths):
            tasks[task_id]["progress"] = f"识别第 {idx+1}/{len(local_paths)} 张图片..."
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(None, ocr_engine.process_image, img_path)
            all_markdown.append(f"<!-- 第 {idx+1} 页 -->\n\n{result['markdown']}")

        combined_md = "\n\n---\n\n".join(all_markdown)

        # 4. 保存本地 Markdown
        md_filename = f"{task_id}_result.md"
        md_path = OUTPUT_DIR / md_filename
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(combined_md)

        # 5. 创建飞书文档回传
        tasks[task_id]["progress"] = "正在创建飞书文档..."
        feishu_url = None
        try:
            # 用源文档的真实标题；取不到再退回通用名
            src_title = await feishu_client.get_doc_title(doc_token)
            title = (f"{src_title}（可编辑版）" if src_title
                     else f"DocLens 转换结果 {datetime.now().strftime('%m-%d %H:%M')}")
            feishu_url = await feishu_client.create_doc_with_content(title, combined_md)
            logger.info(f"飞书文档创建成功: {feishu_url}")
            await feishu_client.append_index_record(title, feishu_url, "飞书链接")
        except Exception as e:
            logger.warning(f"创建飞书文档失败 (结果已保存为本地文件): {e}")

        # 统计
        total_regions = 0
        total_tables = 0
        for img_path in local_paths:
            r = ocr_engine.process_image(img_path)
            total_regions += r["stats"]["region_count"]
            total_tables += r["stats"]["table_count"]

        tasks[task_id]["result"] = {
            "markdown": combined_md,
            "stats": {
                "elapsed": 0,
                "region_count": total_regions,
                "table_count": total_tables,
                "image_count": len(local_paths),
                "page_count": len(local_paths),
            },
            "md_file": md_filename,
            "docx_file": None,
            "feishu_url": feishu_url,
        }
        tasks[task_id]["status"] = "completed"
        tasks[task_id]["progress"] = "完成"

    except Exception as e:
        logger.error(f"飞书任务 {task_id} 失败: {e}", exc_info=True)
        tasks[task_id]["status"] = "failed"
        tasks[task_id]["progress"] = f"失败: {str(e)}"
    finally:
        for p in local_paths:
            try: os.remove(p)
            except OSError: pass


# ---- 公众号/网页文章 -> 飞书 ----

@app.post("/api/convert/article")
async def convert_article(body: dict = Body(...)):
    url = body.get("url", "").strip()
    if not url:
        raise HTTPException(400, "请提供文章链接")
    if not feishu_client.is_configured:
        raise HTTPException(400, "请先配置飞书应用凭证")
    if not feishu_client.has_user_auth:
        raise HTTPException(400, "请先登录飞书账号(右上角『登录飞书』)，以确保生成的文档归你所有")
    if not network_guard.check_feishu_reachable():
        raise HTTPException(400, "连不上飞书服务器，请检查网络")

    task_id = str(uuid.uuid4())[:8]
    tasks[task_id] = {
        "id": task_id,
        "mode": "article",
        "filename": url,
        "status": "processing",
        "progress": "抓取文章...",
        "created_at": datetime.now().isoformat(),
        "result": None,
    }
    asyncio.create_task(_process_article_task(task_id, url))
    return {"task_id": task_id, "status": "processing"}


async def _process_article_task(task_id: str, url: str):
    image_paths = []
    try:
        tasks[task_id]["progress"] = "抓取并解析文章..."
        title, items = await parse_web_article(url)
        image_paths = [p for k, p in items if k == "image"]
        img_n = len(image_paths)
        tbl_n = sum(1 for k, _ in items if k == "table")
        txt_n = sum(1 for k, _ in items if k == "block")

        tasks[task_id]["progress"] = f"写入飞书文档(文字{txt_n}/图{img_n}/表{tbl_n})..."
        feishu_url = await feishu_client.create_doc_from_items(title, items)
        await feishu_client.append_index_record(title, feishu_url, "公众号文章")

        tasks[task_id]["result"] = {
            "markdown": f"# {title}\n\n已转换为飞书文档：{feishu_url}",
            "stats": {"elapsed": 0, "region_count": txt_n,
                      "table_count": tbl_n, "image_count": img_n},
            "md_file": None, "docx_file": None, "feishu_url": feishu_url,
        }
        tasks[task_id]["status"] = "completed"
        tasks[task_id]["progress"] = "完成"
    except Exception as e:
        logger.error(f"文章任务 {task_id} 失败: {e}", exc_info=True)
        tasks[task_id]["status"] = "failed"
        tasks[task_id]["progress"] = f"失败: {str(e)}"
    finally:
        for p in image_paths:
            try: os.remove(p)
            except OSError: pass


# ---- 公众号批量(多个链接/txt) ----

@app.post("/api/convert/articles")
async def convert_articles(body: dict = Body(...)):
    raw = body.get("urls", "")
    if isinstance(raw, str):
        raw = re.split(r"[\s,，;；]+", raw)
    urls = [u.strip() for u in raw if u.strip().startswith("http")]
    if not urls:
        raise HTTPException(400, "没有有效的文章链接")
    if not feishu_client.is_configured:
        raise HTTPException(400, "请先配置飞书应用凭证")
    if not feishu_client.has_user_auth:
        raise HTTPException(400, "请先登录飞书账号(右上角『登录飞书』)，以确保生成的文档归你所有")
    if not network_guard.check_feishu_reachable():
        raise HTTPException(400, "连不上飞书服务器，请检查网络")

    task_id = str(uuid.uuid4())[:8]
    tasks[task_id] = {
        "id": task_id, "mode": "batch", "status": "processing",
        "progress": f"共 {len(urls)} 篇，准备中...",
        "created_at": datetime.now().isoformat(),
        "result": {"items": [{"title": u[:50], "status": "待处理"} for u in urls],
                   "done": 0, "total": len(urls)},
    }
    asyncio.create_task(_process_articles_batch(task_id, urls))
    return {"task_id": task_id, "status": "processing", "total": len(urls)}


async def _process_articles_batch(task_id: str, urls: list):
    items = tasks[task_id]["result"]["items"]
    done = 0
    for i, url in enumerate(urls):
        items[i]["status"] = "处理中"
        tasks[task_id]["progress"] = f"第 {i+1}/{len(urls)} 篇..."
        image_paths = []
        try:
            title, parsed = await parse_web_article(url)
            image_paths = [p for k, p in parsed if k == "image"]
            feishu_url = await feishu_client.create_doc_from_items(title, parsed)
            await feishu_client.append_index_record(title, feishu_url, "公众号文章")
            items[i].update(title=title[:50], status="完成", feishu_url=feishu_url)
        except Exception as e:
            logger.error(f"批量文章 {url} 失败: {e}")
            items[i].update(status="失败", error=str(e))
        finally:
            for p in image_paths:
                try: os.remove(p)
                except OSError: pass
        done += 1
        tasks[task_id]["result"]["done"] = done
        if i < len(urls) - 1:
            await asyncio.sleep(2)  # 每篇间隔，避免公众号限流
    ok = sum(1 for it in items if it["status"] == "完成")
    tasks[task_id]["status"] = "completed"
    tasks[task_id]["progress"] = f"完成：成功 {ok}/{len(urls)} 篇"


# ---- 把已识别结果按需上传到飞书 ----

@app.post("/api/task/{task_id}/to-feishu")
async def task_to_feishu(task_id: str):
    if task_id not in tasks:
        raise HTTPException(404, "任务不存在")
    task = tasks[task_id]
    result = task.get("result")
    if not result or not result.get("markdown"):
        raise HTTPException(400, "该任务还没有识别结果")
    if result.get("feishu_url"):
        return {"feishu_url": result["feishu_url"]}  # 已上传过，直接返回
    if not feishu_client.is_configured:
        raise HTTPException(400, "请先配置飞书应用凭证")
    if not feishu_client.has_user_auth:
        raise HTTPException(400, "请先登录飞书账号(右上角『登录飞书』)，以确保生成的文档归你所有")
    if not network_guard.check_feishu_reachable():
        raise HTTPException(400, "连不上飞书服务器，请检查网络")

    title = task.get("filename") or f"DocLens 识别结果 {datetime.now().strftime('%m-%d %H:%M')}"
    try:
        feishu_url = await feishu_client.create_doc_with_content(title, result["markdown"])
        await feishu_client.append_index_record(title, feishu_url, "本地上传")
    except Exception as e:
        logger.error(f"上传飞书失败: {e}", exc_info=True)
        raise HTTPException(500, f"上传飞书失败: {str(e)}")
    result["feishu_url"] = feishu_url
    return {"feishu_url": feishu_url}


# ---- 通用 API ----

@app.get("/api/task/{task_id}")
async def get_task_status(task_id: str):
    if task_id not in tasks:
        raise HTTPException(404, "任务不存在")
    return tasks[task_id]


def _safe_under(base: Path, filename: str) -> Path:
    """只允许访问 base 目录下的文件，阻断 ../ 路径穿越。"""
    # 只取文件名部分，剥掉任何目录成分
    name = os.path.basename(filename)
    target = (base / name).resolve()
    if base.resolve() not in target.parents and target != base.resolve():
        raise HTTPException(400, "非法文件名")
    return target


@app.get("/api/download/{filename}")
async def download_file(filename: str):
    file_path = _safe_under(OUTPUT_DIR, filename)
    if not file_path.exists():
        raise HTTPException(404, "文件不存在")
    return FileResponse(path=file_path, filename=file_path.name)


@app.get("/api/image/{filename}")
async def get_image(filename: str):
    """MinerU 抽取出的图片(供 Markdown 预览引用)"""
    file_path = _safe_under(OUTPUT_DIR / "images", filename)
    if not file_path.exists():
        raise HTTPException(404, "图片不存在")
    return FileResponse(path=file_path)


@app.get("/api/health")
async def health():
    return {
        "status": "ok",
        "engine": "MinerU",
        "engine_loaded": ocr_engine._initialized,
        "feishu_configured": feishu_client.is_configured,
        "index_url": load_config().get("index_url", ""),
        "network": network_guard.get_status(),
    }


# ============================================================
# 入口
# ============================================================

if __name__ == "__main__":
    import webbrowser

    # 安全：强制只监听本机。除非显式设 DOCLENS_ALLOW_LAN=1 才允许对外(并大声警告)
    bind_host = HOST
    if os.environ.get("DOCLENS_ALLOW_LAN") == "1":
        bind_host = "0.0.0.0"
        print("  " + "!" * 44)
        print("  ⚠️  危险：你开启了对外访问(DOCLENS_ALLOW_LAN=1)")
        print("  ⚠️  本工具没有登录鉴权，同网络的任何人都能用你的飞书身份！")
        print("  ⚠️  仅在你完全清楚后果时使用。")
        print("  " + "!" * 44)
    else:
        bind_host = "127.0.0.1"  # 锁死本机，改 HOST 常量也无法意外暴露

    print()
    print("  ╔══════════════════════════════════════════╗")
    print("  ║              DocLens                      ║")
    print("  ╚══════════════════════════════════════════╝")
    print()
    print(f"  🌐 访问: http://127.0.0.1:{PORT}")
    print(f"  🔒 仅监听本机，外部无法访问" if bind_host == "127.0.0.1" else f"  ⚠️  对外开放中: {bind_host}:{PORT}")
    print(f"  ⏹  Ctrl+C 停止")
    print()

    webbrowser.open(f"http://127.0.0.1:{PORT}")
    uvicorn.run(app, host=bind_host, port=PORT, log_level="warning")
